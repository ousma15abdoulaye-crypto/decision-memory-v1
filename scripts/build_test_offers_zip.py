#!/usr/bin/env python3
"""Génère ``data/test_zip/*.docx`` et ``data/test_zip/test_offers.zip`` (corpus minimal).

À lancer avant commit ou en CI :
    python scripts/build_test_offers_zip.py

Chaque fournisseur a **trois** pièces (offre + NIF + RCCM) avec la même ligne
« SARL » en tête, pour que ``_check_completeness`` du Pass -1 atteigne 1.0 et
évite ``interrupt()`` HITL (sinon LangGraph s'arrête avant ``finalize`` et
``bundle_ids`` reste vide — voir ``src/assembler/graph.py``).
"""

from __future__ import annotations

import sys
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

OUT_DIR = ROOT / "data" / "test_zip"

# (slug fichiers, ligne vendeur unique — doit contenir un token reconnu par
# ``_extract_vendor_name``, ex. SARL, pour le groupement.)
SUPPLIERS: tuple[tuple[str, str], ...] = (
    ("alpha", "FOURNISSEUR ALPHA SARL — corpus test DMS."),
    ("beta", "FOURNISSEUR BETA SARL — corpus test DMS."),
    ("gamma", "FOURNISSEUR GAMMA SARL — corpus test DMS."),
)


def _write_offer(path: Path, vendor_line: str, idx: int) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph(vendor_line)
    doc.add_heading(f"Offre test {idx}", level=1)
    doc.add_paragraph(
        f"Montant total {1000 * idx:,} XOF hors taxes. Délai de livraison 30 jours."
    )
    doc.save(path)


def _write_nif(path: Path, vendor_line: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph(vendor_line)
    doc.add_paragraph(
        "Identification fiscale — NIF 12ABCD56789E. Attestation valide fournisseur."
    )
    doc.save(path)


def _write_rccm(path: Path, vendor_line: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph(vendor_line)
    doc.add_paragraph(
        "Extrait RCCM N° MM-01-A9-12345X — registre du commerce et du crédit mobilier."
    )
    doc.save(path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []

    for i, (slug, vendor_line) in enumerate(SUPPLIERS, start=1):
        trip = (
            (f"{slug}_offre.docx", lambda p: _write_offer(p, vendor_line, i)),
            (f"{slug}_nif.docx", lambda p: _write_nif(p, vendor_line)),
            (f"{slug}_rccm.docx", lambda p: _write_rccm(p, vendor_line)),
        )
        for name, writer in trip:
            p = OUT_DIR / name
            writer(p)
            written.append(p)
            print("wrote", p)

    zpath = OUT_DIR / "test_offers.zip"
    with ZipFile(zpath, "w", compression=ZIP_DEFLATED) as zf:
        for p in written:
            zf.write(p, arcname=p.name)
    print("wrote", zpath)


if __name__ == "__main__":
    main()
