# src/extraction/engine.py
"""ExtractionEngine — M-EXTRACTION-ENGINE
Constitution V3.3.2 §1 (Couche A) + §9 (doctrine échec).
ADR-0002 §2.5 (SLA deux classes).
ADR-0003 §2.2 (ordre exécution).

SLA-A : native_pdf / excel_parser / docx_parser
        → synchrone < 60s
SLA-B : azure / llamaparse / mistral_ocr (cloud-first, retry backoff)
        → asynchrone via extraction_jobs
"""

from __future__ import annotations

# stdlib
import json
import logging
import os
import random
import time
import uuid
from threading import Event

# third-party
import openpyxl  # type: ignore
import pdfplumber  # type: ignore
from docx import Document  # type: ignore

from src.core.api_keys import get_llama_cloud_api_key, get_mistral_api_key

# local
from src.db.connection import get_db_cursor

_logger = logging.getLogger(__name__)
_BACKOFF_WAITER = Event()

# ── Constantes ───────────────────────────────────────────────────

SLA_A_METHODS = {"native_pdf", "excel_parser", "docx_parser"}
SLA_B_METHODS = {"azure", "llamaparse", "mistral_ocr", "tesseract"}
SLA_A_TIMEOUT_S = 60.0
INSUFFICIENT_TEXT_THRESHOLD = 100

STRUCTURED_DATA_EMPTY: dict = {
    "doc_kind": None,
    "language_detected": None,
    "detected_tables": [],
    "detected_sections": [],
    "candidate_criteria": [],
    "candidate_line_items": [],
    "currency_detected": None,
    "dates_detected": [],
    "supplier_candidates": [],
    "_low_confidence": False,
    "_requires_human_review": False,
    "_extraction_duration_ms": None,
    "_sla_class": None,
}


# ── Détection méthode ────────────────────────────────────────────


def detect_method(mime_type: str, file_content: bytes) -> str:
    """Detecte la methode d'extraction selon magic bytes reels.

    Jamais selon Content-Type declare (ADR-0002).
    Cloud-first : scanned PDFs → mistral_ocr (pas tesseract).
    """
    if file_content[:4] == b"%PDF":
        text_markers = (b"/Type /Page", b"BT", b"/Font")
        marker_hits = sum(1 for m in text_markers if m in file_content[:8192])
        if marker_hits >= 2:
            return "native_pdf"
        return "mistral_ocr"

    if file_content[:4] == b"PK\x03\x04":
        if b"xl/" in file_content[:512]:
            return "excel_parser"
        if b"word/" in file_content[:512]:
            return "docx_parser"
        return "mistral_ocr"

    return "mistral_ocr"


# ── Helpers privés ───────────────────────────────────────────────


def get_document(document_id: str) -> dict:
    """
    Charge un document depuis la DB.
    §9 : raise ValueError si introuvable.
    """
    with get_db_cursor() as cur:
        cur.execute(
            """
            SELECT id, case_id, mime_type, storage_uri,
                   extraction_status, extraction_method
            FROM documents
            WHERE id = %s
        """,
            (document_id,),
        )
        doc = cur.fetchone()

    if doc is None:
        raise ValueError(
            f"Document '{document_id}' introuvable. "
            f"Vérifier l'id ou lancer POST /api/cases/{{id}}"
            f"/documents d'abord."
        )
    return dict(doc)


def _update_document_status(document_id: str, status: str) -> None:
    """Met à jour extraction_status sur documents."""
    with get_db_cursor() as cur:
        cur.execute(
            """
            UPDATE documents
            SET extraction_status = %s
            WHERE id = %s
        """,
            (status, document_id),
        )


def _compute_confidence(
    raw_text: str,
    structured: dict,  # noqa: ARG001
) -> float:
    """Heuristique de confiance — valeurs canoniques {0.6, 0.8, 1.0} uniquement.

    §9 : incertitude mesuree, jamais masquee.
    Mapping : <100 chars -> 0.6, <500 chars -> 0.8, >=500 -> 1.0
    La valeur retournee est toujours >= 0.6 ; les branches < 0.6 sont mortes.
    """
    text = raw_text.strip() if raw_text else ""
    if len(text) < INSUFFICIENT_TEXT_THRESHOLD:
        return 0.6
    if len(text) < 500:
        return 0.8
    return 1.0


def _validate_storage_uri(storage_uri: str) -> str:
    """Verifie existence et chemin dans zone autorisee."""
    if not os.path.exists(storage_uri):
        raise FileNotFoundError(f"Document introuvable : {storage_uri}")
    real_uri = os.path.realpath(storage_uri)
    allowed_base = os.environ.get("STORAGE_BASE_PATH", "")
    if not allowed_base:
        import logging as _log

        _log.getLogger(__name__).warning(
            "[OCR] STORAGE_BASE_PATH non defini — utilisation du repertoire courant. "
            "Definir STORAGE_BASE_PATH pour securiser les acces fichier."
        )
        allowed_base = os.getcwd()
    real_base = os.path.realpath(allowed_base)
    if not real_uri.startswith(real_base):
        raise PermissionError(f"Chemin hors zone autorisee : {storage_uri}")
    return real_uri


def _store_extraction(
    document_id: str,
    raw_text: str,
    structured_data: dict,
    method: str,
    confidence: float,
    case_id: str | None = None,
) -> None:
    """Persiste le résultat d'extraction en DB."""
    if case_id is None:
        doc = get_document(document_id)
        case_id = doc.get("case_id")

    extraction_id = f"ext-{uuid.uuid4().hex[:12]}"

    with get_db_cursor() as cur:
        structured_json = json.dumps(structured_data)
        cur.execute(
            """
            INSERT INTO extractions
                (id, case_id, document_id, raw_text, structured_data,
                 extraction_method, confidence_score, extracted_at,
                 data_json, extraction_type, created_at)
            VALUES (%s, %s, %s, %s, %s::jsonb, %s, %s, NOW(),
                    %s, %s, NOW()::TEXT)
        """,
            (
                extraction_id,
                case_id,
                document_id,
                raw_text,
                structured_json,
                method,
                confidence,
                structured_json,  # data_json pour compatibilité
                method,  # extraction_type pour compatibilité
            ),
        )


def _store_error(
    document_id: str,
    job_id: str | None,
    error_code: str,
    error_detail: str,
) -> None:
    """
    §9 : enregistre l'erreur explicitement.
    Jamais silencieux.
    """
    with get_db_cursor() as cur:
        cur.execute(
            """
            INSERT INTO extraction_errors
                (document_id, job_id, error_code, error_message)
            VALUES (%s, %s, %s, %s)
        """,
            (document_id, job_id, error_code, error_detail),
        )


# ── Parseurs ─────────────────────────────────────────────────────


def _extract_native_pdf(storage_uri: str) -> tuple[str, dict]:
    """Extraction PDF natif via pdfplumber."""
    _validate_storage_uri(storage_uri)
    raw_text = ""
    with pdfplumber.open(storage_uri) as pdf:
        for page in pdf.pages:
            raw_text += (page.extract_text() or "") + "\n"

    return raw_text, dict(STRUCTURED_DATA_EMPTY)


def _extract_excel(storage_uri: str) -> tuple[str, dict]:
    """Extraction XLSX via openpyxl."""
    _validate_storage_uri(storage_uri)
    wb = openpyxl.load_workbook(storage_uri, data_only=True)
    try:
        raw_text = ""
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    raw_text += row_text + "\n"
        return raw_text, dict(STRUCTURED_DATA_EMPTY)
    finally:
        wb.close()


def _extract_docx(storage_uri: str) -> tuple[str, dict]:
    """Extraction DOCX via python-docx."""
    _validate_storage_uri(storage_uri)
    doc = Document(storage_uri)
    raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return raw_text, dict(STRUCTURED_DATA_EMPTY)


_OCR_MAX_RETRIES = 3
_OCR_BACKOFF_BASE_S = 2.0
_OCR_JITTER_MAX_S = 0.5

# Exceptions non-retryables : erreurs fatales de configuration ou de fichier.
_NON_RETRYABLE_EXCEPTIONS = (
    ValueError,
    FileNotFoundError,
    ImportError,
    PermissionError,
)


def _retry_cloud_ocr(func, storage_uri: str, method_name: str) -> tuple[str, dict]:
    """Wrapper retry avec backoff exponentiel + jitter pour les extracteurs cloud.

    Exceptions non-retryables (ValueError, FileNotFoundError, ImportError,
    PermissionError) sont re-levées immédiatement sans retry.
    Exceptions réseau (ConnectionError, TimeoutError, OSError) déclenchent
    le backoff exponentiel avec jitter.
    """
    last_exc: Exception | None = None
    for attempt in range(_OCR_MAX_RETRIES):
        t0 = time.monotonic()
        try:
            result = func(storage_uri)
            duration_ms = (time.monotonic() - t0) * 1000
            _logger.info(
                "[OCR] %s OK en %.0fms (tentative %d/%d)",
                method_name,
                duration_ms,
                attempt + 1,
                _OCR_MAX_RETRIES,
            )
            return result
        except _NON_RETRYABLE_EXCEPTIONS as exc:
            _logger.error(
                "[OCR] %s erreur fatale (non-retryable) : %s",
                method_name,
                exc,
            )
            raise
        except Exception as exc:
            last_exc = exc
            duration_ms = (time.monotonic() - t0) * 1000
            if attempt < _OCR_MAX_RETRIES - 1:
                jitter = random.uniform(0, _OCR_JITTER_MAX_S)  # noqa: S311
                wait = _OCR_BACKOFF_BASE_S * (2**attempt) + jitter
                _logger.warning(
                    "[OCR] %s echec tentative %d/%d (%.0fms) : %s — retry dans %.1fs",
                    method_name,
                    attempt + 1,
                    _OCR_MAX_RETRIES,
                    duration_ms,
                    exc,
                    wait,
                )
                _BACKOFF_WAITER.wait(wait)
            else:
                _logger.error(
                    "[OCR] %s echec definitif apres %d tentatives (%.0fms) : %s",
                    method_name,
                    _OCR_MAX_RETRIES,
                    duration_ms,
                    exc,
                )
    raise last_exc  # type: ignore[misc]


def _extract_llamaparse(storage_uri: str) -> tuple[str, dict]:
    """Extraction via LlamaParse (SLA-B).

    La clé API est lue via get_llama_cloud_api_key() : LLAMADMS puis
    LLAMA_CLOUD_API_KEY (jamais dans le code).
    Lève APIKeyMissingError si les deux sont absentes.
    """
    _validate_storage_uri(storage_uri)
    api_key = get_llama_cloud_api_key()  # raises APIKeyMissingError if absent

    try:
        from llama_parse import LlamaParse  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "llama-parse n'est pas installé. "
            "Ajouter 'llama-parse' à requirements.txt."
        ) from exc

    parser = LlamaParse(api_key=api_key, result_type="markdown")
    documents = parser.load_data(storage_uri)
    raw_text = "\n".join(doc.text for doc in documents)
    return raw_text, dict(STRUCTURED_DATA_EMPTY)


# Mandat 2 — Files API : stream upload, zéro RAM, cleanup garanti
_MISTRAL_OCR_MAX_BYTES = 50 * 1024 * 1024  # 50 MB guard-rail
_MISTRAL_OCR_SUPPORTED_MIMES = ("image/", "application/pdf")


def _ocr_pages_to_text(response: object) -> str:
    """Extrait le texte Markdown depuis response.pages (mistral-ocr-latest)."""
    if hasattr(response, "pages") and response.pages:
        return "\n\n".join(
            getattr(p, "markdown", "") or getattr(p, "text", "") for p in response.pages
        )
    return getattr(response, "text", "") or ""


def _mistral_client_factory():
    """Retourne la classe constructeur Mistral(api_key=...) selon la version du SDK.

    mistralai récent : ``Mistral`` vit souvent sous ``mistralai.client`` ; un
    ``ImportError`` sur ``from mistralai import Mistral`` ne signifie pas que le
    package est absent (sous-import ou API déplacée).
    """
    try:
        from mistralai import Mistral as MistralCls  # type: ignore
    except ImportError as exc_top:
        try:
            from mistralai.client import Mistral as MistralCls  # type: ignore
        except ImportError as exc_client:
            raise ImportError(
                "Impossible d'importer la classe Mistral depuis mistralai "
                f"(toplevel: {exc_top!r}; mistralai.client: {exc_client!r}). "
                "Vérifiez mistralai (requirements.txt) et l'installation du package."
            ) from exc_client
    return MistralCls


def _ensure_ssl_env() -> None:
    """Configure SSL_CERT_FILE / REQUESTS_CA_BUNDLE si absents.

    Sur les postes avec proxy d'entreprise, le CA store système peut bloquer les
    appels HTTPS vers api.mistral.ai. Certifi fournit un bundle Mozilla connu.
    Positionne les variables d'environnement uniquement si elles ne sont pas déjà
    définies — respecte toute configuration manuelle du système.
    """
    if os.environ.get("SSL_CERT_FILE") and os.environ.get("REQUESTS_CA_BUNDLE"):
        return
    try:
        import certifi  # type: ignore

        bundle = certifi.where()
        os.environ.setdefault("SSL_CERT_FILE", bundle)
        os.environ.setdefault("REQUESTS_CA_BUNDLE", bundle)
    except ImportError:
        pass  # certifi optionnel — pas de crash si absent


def _extract_mistral_ocr(storage_uri: str) -> tuple[str, dict]:
    """OCR via Mistral Files API (SLA-B) — Mandat 2 / ADR-M11-002.

    Pattern : upload (stream) → ocr.process(file_id) → delete (finally).
    RAM consommée : ~2 MB (chunk réseau) au lieu de 117 MB (base64 full-load).
    Fallback Azure actif si AZURE_FORM_RECOGNIZER_ENDPOINT défini.
    Lève APIKeyMissingError si MISTRAL_API_KEY absent.
    SSL/TLS : certifi utilisé si SSL_CERT_FILE absent (proxy d'entreprise).
    """
    _validate_storage_uri(storage_uri)
    api_key = get_mistral_api_key()

    # Resilience SSL/TLS pour les postes avec proxy d'entreprise
    _ensure_ssl_env()

    # Guard taille avant ouverture
    file_size = os.path.getsize(storage_uri)
    if file_size > _MISTRAL_OCR_MAX_BYTES:
        raise ValueError(
            f"Fichier trop volumineux ({file_size / 1024 / 1024:.1f} MB > "
            f"{_MISTRAL_OCR_MAX_BYTES / 1024 / 1024:.0f} MB). "
            "Découper le PDF avant envoi."
        )

    # Détection MIME sans charger le fichier entier en RAM
    mime = _detect_mime_from_header(storage_uri)
    azure_endpoint = os.environ.get("AZURE_FORM_RECOGNIZER_ENDPOINT", "")

    if not any(mime.startswith(p) for p in _MISTRAL_OCR_SUPPORTED_MIMES):
        if azure_endpoint:
            return _extract_azure_ocr_fallback(storage_uri, mime)
        raise ValueError(
            f"Mistral OCR : type non supporté '{mime}'. "
            "Configurer AZURE_FORM_RECOGNIZER_ENDPOINT pour le fallback."
        )

    from src.couche_a.llm_router import TIER_1_OCR_MODEL

    mistral_cls = _mistral_client_factory()
    client = mistral_cls(api_key=api_key)
    file_name = os.path.basename(storage_uri)
    uploaded_id: str | None = None

    try:
        # Étape 1 — upload en streaming (file handle, pas de bytes en mémoire)
        with open(storage_uri, "rb") as fh:
            uploaded = client.files.upload(
                file={"file_name": file_name, "content": fh},
                purpose="ocr",
            )
        uploaded_id = uploaded.id

        # Étape 2 — OCR via file_id (URL signée gérée côté Mistral Cloud)
        response = client.ocr.process(
            model=TIER_1_OCR_MODEL,
            document={"type": "file_id", "file_id": uploaded_id},
        )
        raw_text = _ocr_pages_to_text(response)

    finally:
        # Étape 3 — cleanup obligatoire quelle que soit l'issue
        if uploaded_id:
            try:
                client.files.delete(file_id=uploaded_id)
            except Exception as cleanup_err:  # noqa: BLE001
                import logging as _log

                _log.getLogger(__name__).warning(
                    "[OCR] Fichier Mistral Cloud non supprimé id=%s : %s",
                    uploaded_id,
                    cleanup_err,
                )

    return raw_text, dict(STRUCTURED_DATA_EMPTY)


def _detect_mime_from_header(storage_uri: str) -> str:
    """Détecte le MIME en lisant seulement les 512 premiers octets (magic bytes).

    Si filetype retourne application/octet-stream mais que les magic bytes sont
    %PDF, force application/pdf — évite le rejet silencieux par Mistral OCR.
    """
    try:
        import filetype as _ft  # type: ignore

        with open(storage_uri, "rb") as fh:
            header = fh.read(512)
        detected = _ft.guess(header)
        mime = detected.mime if detected else "application/octet-stream"
        # PDF magic bytes (%PDF) non reconnu par filetype → forcer application/pdf
        if mime == "application/octet-stream" and header[:4] == b"%PDF":
            return "application/pdf"
        return mime
    except Exception:
        # Dernier recours : lire les magic bytes directement
        try:
            with open(storage_uri, "rb") as fh:
                if fh.read(4) == b"%PDF":
                    return "application/pdf"
        except Exception:
            pass
        return "application/octet-stream"


def _extract_azure_ocr_fallback(storage_uri: str, mime: str) -> tuple[str, dict]:
    """Fallback Azure Form Recognizer — stream, pas de bytes en RAM."""
    endpoint = os.environ.get("AZURE_FORM_RECOGNIZER_ENDPOINT", "")
    key = os.environ.get("AZURE_FORM_RECOGNIZER_KEY", "")
    if not endpoint or not key:
        raise RuntimeError(
            "[OCR] Azure fallback non configuré "
            "(AZURE_FORM_RECOGNIZER_ENDPOINT ou AZURE_FORM_RECOGNIZER_KEY manquant)."
        )
    try:
        from azure.ai.formrecognizer import DocumentAnalysisClient  # type: ignore
        from azure.core.credentials import AzureKeyCredential  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "azure-ai-formrecognizer non installé. "
            "Ajouter 'azure-ai-formrecognizer' à requirements.txt."
        ) from exc

    az_client = DocumentAnalysisClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )
    with open(storage_uri, "rb") as fh:
        poller = az_client.begin_analyze_document("prebuilt-read", fh)
    result = poller.result()
    raw_text = (
        "\n".join(p.content for p in result.paragraphs) if result.paragraphs else ""
    )
    return raw_text, dict(STRUCTURED_DATA_EMPTY)


def _dispatch_extraction(
    doc: dict,
    method: str,
) -> tuple[str, dict]:
    """Routing vers le parseur selon la méthode.

    SLA-A (sync) : native_pdf, excel_parser, docx_parser.
    SLA-B (async, cloud-first) : mistral_ocr, llamaparse, azure.
    Online-first — pas de tesseract local (hors scope cloud-first).
    """
    if method == "native_pdf":
        return _extract_native_pdf(doc["storage_uri"])
    if method == "excel_parser":
        return _extract_excel(doc["storage_uri"])
    if method == "docx_parser":
        return _extract_docx(doc["storage_uri"])
    # SLA-B : cloud OCR (online-first) avec retry backoff exponentiel
    # Compat legacy: "tesseract" est maintenu comme alias DB/API vers mistral_ocr.
    if method == "tesseract":
        return _retry_cloud_ocr(_extract_mistral_ocr, doc["storage_uri"], "mistral_ocr")
    if method == "mistral_ocr":
        return _retry_cloud_ocr(_extract_mistral_ocr, doc["storage_uri"], "mistral_ocr")
    if method == "llamaparse":
        return _retry_cloud_ocr(_extract_llamaparse, doc["storage_uri"], "llamaparse")
    if method == "azure":
        return _retry_cloud_ocr(
            lambda uri: _extract_azure_ocr_fallback(uri, "application/pdf"),
            doc["storage_uri"],
            "azure",
        )
    raise ValueError(
        f"Méthode inconnue pour dispatch : '{method}'. "
        f"SLA-A : {SLA_A_METHODS} | SLA-B : {SLA_B_METHODS}"
    )


# ── Fonctions publiques ──────────────────────────────────────────


def extract_sync(document_id: str) -> dict:
    """
    SLA-A : extraction synchrone < 60s.
    Méthodes : native_pdf, excel_parser, docx_parser.
    §9 : échec explicite si timeout ou erreur.
    """
    start_ms = time.monotonic() * 1000

    doc = get_document(document_id)
    method = doc.get("extraction_method") or "native_pdf"

    if method not in SLA_A_METHODS:
        raise ValueError(
            f"extract_sync appelé avec méthode SLA-B '{method}'. "
            f"Utiliser extract_async pour {SLA_B_METHODS}."
        )

    _update_document_status(document_id, "processing")

    try:
        raw_text, structured_data = _dispatch_extraction(doc, method)
        duration_ms = (time.monotonic() * 1000) - start_ms

        # §9 : SLA-A violation → échec explicite
        if duration_ms > SLA_A_TIMEOUT_S * 1000:
            _store_error(
                document_id,
                None,
                "TIMEOUT_SLA_A",
                f"SLA-A violé : {duration_ms:.0f}ms > 60000ms",
            )
            _update_document_status(document_id, "failed")
            raise TimeoutError(
                f"SLA-A violé : {duration_ms:.0f}ms. Document {document_id}."
            )

        confidence = _compute_confidence(raw_text, structured_data)

        structured_data["_extraction_duration_ms"] = duration_ms
        structured_data["_sla_class"] = "A"

        _store_extraction(
            document_id,
            raw_text,
            structured_data,
            method,
            confidence,
            case_id=doc.get("case_id"),
        )
        _update_document_status(document_id, "done")

        return {
            "document_id": document_id,
            "status": "done",
            "method": method,
            "sla_class": "A",
            "duration_ms": duration_ms,
            "confidence": confidence,
            "requires_human_review": structured_data["_requires_human_review"],
        }

    except TimeoutError:
        raise
    except Exception as exc:
        # §9 : échec explicite, jamais silencieux
        _update_document_status(document_id, "failed")
        _store_error(
            document_id,
            None,
            "PARSE_ERROR",
            str(exc),
        )
        raise


def extract_async(document_id: str, method: str) -> dict:
    """
    SLA-B : mise en queue OCR asynchrone.
    Retourne immédiatement avec job_id + status pending.
    """
    if method not in SLA_B_METHODS:
        raise ValueError(
            f"extract_async appelé avec méthode SLA-A '{method}'. "
            f"Utiliser extract_sync pour {SLA_A_METHODS}."
        )

    with get_db_cursor() as cur:
        cur.execute(
            """
            INSERT INTO extraction_jobs
                (document_id, method, sla_class, status)
            VALUES (%s, %s, 'B', 'pending')
            RETURNING id
        """,
            (document_id, method),
        )
        job = cur.fetchone()

    return {
        "document_id": document_id,
        "job_id": str(job["id"]),
        "status": "pending",
        "method": method,
        "sla_class": "B",
        "message": (
            "Extraction OCR mise en queue. "
            "Consulter GET /api/extractions/jobs/"
            f"{job['id']}/status"
        ),
    }


def process_extraction_job(job_id: str) -> dict:
    """Execute a queued SLA-B extraction job.

    Called by the async worker when it picks a job from the queue.
    Transitions: pending → processing → done | failed.
    §9 : échec explicite, jamais silencieux.
    """
    with get_db_cursor() as cur:
        cur.execute(
            """
            SELECT ej.id, ej.document_id, ej.method, ej.status,
                   d.storage_uri, d.case_id
            FROM extraction_jobs ej
            JOIN documents d ON d.id = ej.document_id
            WHERE ej.id = %s
        """,
            (job_id,),
        )
        job = cur.fetchone()

    if job is None:
        raise ValueError(f"Job '{job_id}' introuvable.")

    if job["status"] != "pending":
        raise ValueError(
            f"Job '{job_id}' n'est pas en attente (status={job['status']}). "
            f"Seuls les jobs 'pending' peuvent être traités."
        )

    method = job["method"]
    document_id = job["document_id"]

    if method not in SLA_B_METHODS:
        raise ValueError(
            f"Méthode inconnue pour SLA-B : '{method}'. "
            f"Méthodes valides : {SLA_B_METHODS}."
        )

    # pending → processing
    with get_db_cursor() as cur:
        cur.execute(
            "UPDATE extraction_jobs SET status = 'processing', "
            "started_at = NOW() WHERE id = %s",
            (job_id,),
        )
    _update_document_status(document_id, "processing")

    start_ms = time.monotonic() * 1000

    try:
        doc = {"storage_uri": job["storage_uri"]}
        raw_text, structured_data = _dispatch_extraction(doc, method)
        duration_ms = (time.monotonic() * 1000) - start_ms

        confidence = _compute_confidence(raw_text, structured_data)

        structured_data["_extraction_duration_ms"] = duration_ms
        structured_data["_sla_class"] = "B"

        _store_extraction(
            document_id,
            raw_text,
            structured_data,
            method,
            confidence,
            case_id=job.get("case_id"),
        )

        # processing → done
        with get_db_cursor() as cur:
            cur.execute(
                "UPDATE extraction_jobs SET status = 'done', "
                "completed_at = NOW(), duration_ms = %s WHERE id = %s",
                (duration_ms, job_id),
            )
        _update_document_status(document_id, "done")

        return {
            "job_id": job_id,
            "document_id": document_id,
            "status": "done",
            "method": method,
            "sla_class": "B",
            "duration_ms": duration_ms,
            "confidence": confidence,
        }

    except Exception as exc:
        # processing → failed
        with get_db_cursor() as cur:
            cur.execute(
                "UPDATE extraction_jobs SET status = 'failed', "
                "completed_at = NOW(), error_message = %s WHERE id = %s",
                (str(exc)[:500], job_id),
            )
        _update_document_status(document_id, "failed")
        _store_error(document_id, job_id, "PARSE_ERROR", str(exc))
        raise
