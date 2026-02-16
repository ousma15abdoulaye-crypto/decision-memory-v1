"""
PV Word Generator – Ouverture et Analyse
Constitution V2.1 § 4.1 – Manuel SCI SC-PR-02 § 4.3.2, §5.3
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict


def _format_date(val):
    """Format date/datetime pour affichage."""
    if val is None:
        return ""
    if hasattr(val, "strftime"):
        return val.strftime("%d/%m/%Y %H:%M")
    return str(val)


def generate_pv_ouverture(case_data: Dict, output_dir: Path) -> Path:
    """
    Génère le PV d'ouverture des offres (format Word).

    Sections (Manuel §4.3.2):
    1. Informations générales
    2. Membres comité ouverture
    3. Soumissionnaires invités
    4. Registre soumissions reçues
    5. Échantillons reçus
    6. Observations générales
    7. Délégation analyse technique
    8. Signatures
    """
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Titre principal
    title = doc.add_heading("PROCÈS-VERBAL D'OUVERTURE DES OFFRES", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x36, 0x60, 0x92)

    doc.add_paragraph(f"Référence RFQ : {case_data.get('dao_ref', '')}")
    doc.add_paragraph(f"Objet du marché : {case_data.get('market_title', '')}")
    doc.add_paragraph()

    # Section 1: Informations générales
    doc.add_heading("1. INFORMATIONS GÉNÉRALES", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Grid Accent 1"

    infos = [
        ("Objet du marché", case_data.get("market_title", "")),
        ("Référence RFQ", case_data.get("dao_ref", "")),
        ("Date limite soumission", _format_date(case_data.get("deadline"))),
        ("Date ouverture", _format_date(case_data.get("opening_date", datetime.now()))),
        ("Lieu", case_data.get("location", "Bureau Procurement")),
    ]
    for i, (label, val) in enumerate(infos):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(val)

    # Section 2: Membres comité ouverture
    doc.add_heading("2. MEMBRES COMITÉ D'OUVERTURE", level=1)
    doc.add_paragraph("Membres présents (minimum 2) :")
    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Nom/Prénom", "Fonction", "Signature"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for i in range(1, 3):
        for j in range(3):
            table.rows[i].cells[j].text = "_______________________"

    # --- Échantillons reçus (amendement A10) ---
    doc.add_heading("Échantillons reçus", level=2)
    p = doc.add_paragraph()
    p.add_run("Soumis : ").bold = True
    p.add_run("☐ Oui  ☐ Non\n")
    p.add_run("Description : ").bold = True
    p.add_run("_" * 50)

    # --- Délégation technique (amendement A10) ---
    doc.add_heading("Délégation technique", level=2)
    p = doc.add_paragraph()
    p.add_run("Présence : ").bold = True
    p.add_run("☐ Oui  ☐ Non\n")
    p.add_run("Noms / Fonctions : ").bold = True
    p.add_run("_" * 50)

    # Section 3: Soumissionnaires invités
    doc.add_heading("3. SOUMISSIONNAIRES INVITÉS", level=1)
    invited = case_data.get("invited_suppliers", [])
    doc.add_paragraph(f"Total : {len(invited)} soumissionnaires")
    for inv in invited:
        doc.add_paragraph(f"• {inv}", style="List Bullet")

    # Section 4: Registre soumissions reçues
    doc.add_heading("4. REGISTRE DES SOUMISSIONS REÇUES", level=1)
    suppliers = case_data.get("suppliers", [])
    table = doc.add_table(rows=len(suppliers) + 1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["N°", "Soumissionnaire", "Mode", "Date", "Heure", "Conformité"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    suppliers_sorted = sorted(
        suppliers, key=lambda x: x.get("timestamp", datetime.now())
    )
    for idx, sup in enumerate(suppliers_sorted, start=1):
        row = table.rows[idx]
        row.cells[0].text = str(idx)
        row.cells[1].text = sup.get("supplier_name", "")
        row.cells[2].text = sup.get("mode_depot", "")
        ts = sup.get("timestamp", datetime.now())
        row.cells[3].text = ts.strftime("%d/%m/%Y")
        row.cells[4].text = ts.strftime("%H:%M:%S")
        row.cells[5].text = "Conforme"

    # Section 5: Échantillons reçus
    doc.add_heading("5. ÉCHANTILLONS REÇUS", level=1)
    samples = [s for s in suppliers_sorted if s.get("has_sample")]
    if samples:
        for s in samples:
            doc.add_paragraph(f"• {s.get('supplier_name')}", style="List Bullet")
    else:
        doc.add_paragraph("Aucun échantillon reçu.")

    # Section 6: Observations générales
    doc.add_heading("6. OBSERVATIONS GÉNÉRALES", level=1)
    doc.add_paragraph("\n\n\n")

    # Section 7: Délégation analyse technique
    doc.add_heading("7. DÉLÉGATION ANALYSE TECHNIQUE", level=1)
    doc.add_paragraph(
        "Le comité a décidé de déléguer l'analyse des critères techniques suivants :"
    )
    doc.add_paragraph("• _________________________________________________")
    doc.add_paragraph("• _________________________________________________")

    # Section 8: Signatures
    doc.add_heading("8. SIGNATURES", level=1)
    table = doc.add_table(rows=3, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Nom", "Fonction", "Date", "Signature"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for i in range(1, 3):
        for j in range(4):
            table.rows[i].cells[j].text = "_________________"

    # Footer
    doc.add_page_break()
    doc.add_paragraph(
        f"Fait à ________________, le {datetime.now().strftime('%d/%m/%Y')}"
    )

    # --- Échantillons reçus (amendement A10) ---
    doc.add_heading("Échantillons reçus", level=2)

    table = doc.add_table(rows=2, cols=3)
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Fournisseur"
    hdr_cells[1].text = "Description échantillon"
    hdr_cells[2].text = "Conforme"

    row_cells = table.rows[1].cells
    row_cells[0].text = "_" * 20
    row_cells[1].text = "_" * 40
    row_cells[2].text = "☐ Oui  ☐ Non"

    doc.add_paragraph()

    # --- Délégation technique (amendement A10) ---
    doc.add_heading("Délégation de pouvoir technique", level=2)

    p = doc.add_paragraph()
    p.add_run("Le comité délègue l'évaluation technique à : ").bold = True
    p.add_run("_" * 40)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Fonction : ").bold = True
    p.add_run("_" * 40)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Date de la délégation : ").bold = True
    p.add_run("_" * 20)

    filename = f"PV_Ouverture_{case_data.get('case_reference', 'case')}.docx"
    filepath = output_dir / filename
    doc.save(filepath)
    return filepath


def generate_pv_analyse(case_data: Dict, cba_summary: Dict, output_dir: Path) -> Path:
    """
    Génère le PV d'analyse des offres (format Word).

    Sections (Manuel §5.3):
    1. Référence PV ouverture
    2. Critères évaluation
    3. Synthèse notation
    4. Négociation (si applicable)
    5. Recommandation comité
    6. Justification décision
    7. Annexes
    8. Signatures
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Titre
    title = doc.add_heading("PROCÈS-VERBAL D'ANALYSE DES OFFRES", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x36, 0x60, 0x92)

    doc.add_paragraph(f"Référence RFQ : {case_data.get('dao_ref', '')}")
    doc.add_paragraph()

    # Section 1: Référence PV ouverture
    doc.add_heading("1. RÉFÉRENCE PV OUVERTURE", level=1)
    opening_date = case_data.get("opening_date", datetime.now())
    doc.add_paragraph(f"Date du PV d'ouverture : {_format_date(opening_date)}")
    doc.add_paragraph(
        f"Référence : PV_Ouverture_{case_data.get('case_reference', 'case')}.docx"
    )

    # Section 2: Critères évaluation
    doc.add_heading("2. CRITÈRES D'ÉVALUATION", level=1)

    # 2.1 Critères essentiels (Pass/Fail)
    doc.add_heading("2.1 Critères Essentiels (Pass/Fail)", level=2)
    suppliers = case_data.get("suppliers", [])
    table = doc.add_table(rows=len(suppliers) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Fournisseur", "Résultat", "Justification"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for idx, sup in enumerate(suppliers, start=1):
        row = table.rows[idx]
        row.cells[0].text = sup.get("supplier_name", "")
        row.cells[1].text = "Pass"
        row.cells[2].text = ""

    # 2.2 Capacité Technique
    doc.add_heading("2.2 Capacité Technique (≥50%)", level=2)

    # Section 3: Synthèse notation
    doc.add_heading("3. SYNTHÈSE NOTATION", level=1)
    classement = cba_summary.get("classement", [])
    table = doc.add_table(rows=len(classement) + 1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = [
        "Classement",
        "Fournisseur",
        "Technique",
        "Durabilité",
        "Commercial",
        "Note Finale",
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for idx, entry in enumerate(classement, start=1):
        row = table.rows[idx]
        row.cells[0].text = str(idx)
        row.cells[1].text = entry.get("supplier_name", "")
        row.cells[2].text = f"{entry.get('technical_score', 0):.2f}"
        row.cells[3].text = f"{entry.get('sustainability_score', 0):.2f}"
        row.cells[4].text = f"{entry.get('commercial_score', 0):.2f}"
        row.cells[5].text = f"{entry.get('final_score', 0):.2f}"

    # --- Section Négociation (amendement A11) ---
    doc.add_page_break()
    doc.add_heading("Négociation", level=1)
    doc.add_heading("Offres retenues pour négociation", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Fournisseur"
    hdr[1].text = "Lot"
    hdr[2].text = "Montant initial"
    hdr[3].text = "Commentaire"
    # Laisse 5 lignes vides
    for _ in range(5):
        cells = table.add_row().cells
        for c in cells:
            c.text = "_" * 15

    doc.add_heading("Résultat de la négociation", level=2)
    p = doc.add_paragraph()
    p.add_run("Nouveau montant : ").bold = True
    p.add_run("_" * 30 + "\n")
    p.add_run("Délai ajusté : ").bold = True
    p.add_run("_" * 30 + "\n")
    p.add_run("Conditions particulières : ").bold = True
    p.add_run("_" * 50)

    # --- Revue Procurement (amendement A11) ---
    doc.add_heading("Revue Procurement", level=2)
    p = doc.add_paragraph()
    p.add_run("Conformité technique : ").bold = True
    p.add_run("☐ Conforme  ☐ Non conforme  ☐ Partiellement conforme\n")
    p.add_run("Observations : ").bold = True
    p.add_run("_" * 50)

    # --- Validation Head of Supply Chain (amendement A11) ---
    doc.add_heading("Validation", level=2)
    p = doc.add_paragraph()
    p.add_run("Nom du Head of Supply Chain : ").bold = True
    p.add_run("_" * 40 + "\n")
    p.add_run("Date : ").bold = True
    p.add_run("_" * 20 + "\n")
    p.add_run("Signature : ").bold = True
    p.add_run("_" * 30)

    # Footer
    doc.add_page_break()
    doc.add_paragraph(
        f"Fait à ________________, le {datetime.now().strftime('%d/%m/%Y')}"
    )

    doc.add_page_break()

    # --- Négociation (amendement A11) ---
    doc.add_heading("Procès-verbal de négociation", level=2)

    doc.add_heading("Offres retenues pour négociation", level=3)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Fournisseur"
    hdr_cells[1].text = "Lot"
    hdr_cells[2].text = "Montant initial (USD)"
    hdr_cells[3].text = "Points négociables"

    for i in range(1, 4):
        for j in range(4):
            table.rows[i].cells[j].text = "_" * 15

    doc.add_paragraph()

    doc.add_heading("Résultat de la négociation", level=3)

    table2 = doc.add_table(rows=2, cols=4)
    table2.style = "Light Grid Accent 1"

    hdr2_cells = table2.rows[0].cells
    hdr2_cells[0].text = "Fournisseur"
    hdr2_cells[1].text = "Montant final (USD)"
    hdr2_cells[2].text = "Délai ajusté"
    hdr2_cells[3].text = "Conditions particulières"

    for j in range(4):
        table2.rows[1].cells[j].text = "_" * 15

    doc.add_paragraph()

    # --- Revue Procurement (amendement A11) ---
    doc.add_heading("Revue Procurement", level=2)

    p = doc.add_paragraph()
    p.add_run("Conformité technique : ").bold = True
    p.add_run("☐ Conforme  ☐ Non conforme  ☐ Partiellement conforme")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Observations : ").bold = True
    doc.add_paragraph("_" * 80)
    doc.add_paragraph("_" * 80)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Révisé par : ").bold = True
    p.add_run("_" * 40)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Date : ").bold = True
    p.add_run("_" * 20)

    doc.add_paragraph()

    # --- Validation finale (amendement A11) ---
    doc.add_heading("Validation finale", level=2)

    p = doc.add_paragraph()
    p.add_run("Nom du Head of Supply Chain : ").bold = True
    p.add_run("_" * 40)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Date de validation : ").bold = True
    p.add_run("_" * 20)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Signature : ").bold = True
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_" * 30)

    filename = f"PV_Analyse_{case_data.get('case_reference', 'case')}.docx"
    filepath = output_dir / filename
    doc.save(filepath)
    return filepath
