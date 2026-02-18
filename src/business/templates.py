"""
Template generation functions for CBA and PV documents.
Handles template-adaptive generation for procurement analysis.
"""

import json
import re
import uuid
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

from src.core.config import OUTPUTS_DIR
from src.core.models import CBATemplateSchema, DAOCriterion
from src.db import db_execute, get_connection


def analyze_cba_template(template_path: str) -> CBATemplateSchema:
    """Analyse dynamique structure template CBA"""
    wb = load_workbook(template_path)
    ws = wb.active

    supplier_header_row = None
    supplier_name_row = None
    supplier_cols: list[int] = []

    # Detect supplier header
    for row_idx in range(1, min(12, ws.max_row + 1)):
        for col_idx in range(1, ws.max_column + 1):
            val = ws.cell(row_idx, col_idx).value
            if val and isinstance(val, str):
                if re.search(r"(?i)(soumissionnaire|supplier|fournisseur|bidder)", val):
                    supplier_header_row = row_idx
                    # Detect supplier columns
                    for c in range(col_idx, min(col_idx + 15, ws.max_column + 1)):
                        if ws.cell(row_idx, c).value:
                            supplier_cols.append(c)
                    break
        if supplier_header_row:
            break

    supplier_name_row = supplier_header_row + 1 if supplier_header_row else None

    # Detect criteria rows
    criteria_start_row = None
    criteria_rows: list[dict[str, Any]] = []

    if supplier_name_row:
        for row_idx in range(
            supplier_name_row + 1, min(supplier_name_row + 40, ws.max_row + 1)
        ):
            _ = ws.cell(row_idx, 1).value or ""
            col_b = ws.cell(row_idx, 2).value or ""

            if isinstance(col_b, str) and len(col_b) > 5:
                if not criteria_start_row:
                    criteria_start_row = row_idx

                # Type detection
                ctype = "autre"
                if re.search(r"(?i)(essent|mandatory|obligatoire)", col_b):
                    ctype = "essentiels"
                elif re.search(r"(?i)(technique|technical|capacity)", col_b):
                    ctype = "technique"
                elif re.search(r"(?i)(commercial|price|prix|co[uû]t)", col_b):
                    ctype = "commercial"
                elif re.search(r"(?i)(durabilit[ée]|sustainability)", col_b):
                    ctype = "durabilite"

                criteria_rows.append(
                    {"row": row_idx, "name": col_b[:150], "type": ctype}
                )

    schema = CBATemplateSchema(
        template_id=str(uuid.uuid4()),
        template_name=Path(template_path).name,
        supplier_header_row=supplier_header_row or 0,
        supplier_name_row=supplier_name_row or 0,
        supplier_cols=supplier_cols,
        criteria_start_row=criteria_start_row or 0,
        criteria_rows=criteria_rows,
        sheets=wb.sheetnames,
        meta={"detected_at": datetime.utcnow().isoformat()},
    )

    return schema


def fill_cba_adaptive(
    template_path: str,
    case_id: str,
    suppliers: list[dict],
    dao_criteria: list[DAOCriterion],
) -> str:
    """
    Remplissage adaptatif basé sur structure template détectée.
    1 DAO = 1 template spécifique = 1 CBA unique.

    COMPORTEMENT ATTENDU:
    - Noms fournisseurs RÉELS (pas d'IDs, pas de hash)
    - Données manquantes → "REVUE MANUELLE" avec surlignage ORANGE
    - Aucun onglet debug visible dans l'export final
    - Aucune note magique, aucune élimination implicite
    """
    schema = analyze_cba_template(template_path)

    # Save schema to DB for future learning
    with get_connection() as conn:
        db_execute(
            conn,
            """
            INSERT INTO cba_template_schemas (id, case_id, template_name, structure_json, created_at)
            VALUES (:tid, :cid, :tname, :struct, :ts)
        """,
            {
                "tid": schema.template_id,
                "cid": case_id,
                "tname": schema.template_name,
                "struct": json.dumps(asdict(schema), ensure_ascii=False),
                "ts": datetime.utcnow().isoformat(),
            },
        )

    # Load template
    wb = load_workbook(template_path)
    ws = wb.active

    # Couleur ORANGE pour REVUE MANUELLE (spec conforme)
    ORANGE_FILL = PatternFill(
        start_color="FFC000", end_color="FFC000", fill_type="solid"
    )
    REVUE_MANUELLE = "REVUE MANUELLE"

    # Metadata section (top)
    ws["A1"] = "Decision Memory System — CBA Adaptatif"
    ws["A2"] = f"Case ID: {case_id}"
    ws["A3"] = f"Generated: {datetime.utcnow().isoformat()}"
    ws["A4"] = f"Template: {schema.template_name}"

    # Fill supplier names (detected row/cols) - NOMS RÉELS uniquement
    if schema.supplier_name_row and schema.supplier_cols:
        for idx, supplier in enumerate(suppliers[: len(schema.supplier_cols)]):
            col = schema.supplier_cols[idx]
            supplier_name = supplier.get("supplier_name", "")

            # Vérifier que ce n'est pas un ID
            if not supplier_name or supplier_name in [
                "SUPPLIER_UNKNOWN",
                "FOURNISSEUR_INCONNU",
            ]:
                supplier_name = REVUE_MANUELLE

            # Écriture unique de la cellule
            cell = ws.cell(schema.supplier_name_row, col)
            cell.value = supplier_name
            if supplier_name == REVUE_MANUELLE:
                cell.fill = ORANGE_FILL

    # Fill commercial data (detected criteria rows)
    for row_idx in range(
        schema.criteria_start_row, min(schema.criteria_start_row + 50, ws.max_row + 1)
    ):
        label = ws.cell(row_idx, 2).value or ""

        for idx, supplier in enumerate(suppliers[: len(schema.supplier_cols)]):
            col = schema.supplier_cols[idx]
            cell = ws.cell(row_idx, col)

            # Vérifier le package_status du fournisseur
            _ = supplier.get("package_status", "UNKNOWN")
            has_financial = supplier.get("has_financial", False)

            # Mapping guided by criteria
            if re.search(r"(?i)(prix|price|montant|co[uû]t)", label):
                if has_financial:
                    val = supplier.get("total_price")
                    if val:
                        cell.value = val
                        if supplier.get("total_price_source"):
                            from openpyxl.comments import Comment

                            cell.comment = Comment(
                                supplier["total_price_source"], "DMS"
                            )
                    else:
                        cell.value = REVUE_MANUELLE
                        cell.fill = ORANGE_FILL
                else:
                    # Offre sans partie financière
                    cell.value = REVUE_MANUELLE
                    cell.fill = ORANGE_FILL

            elif re.search(r"(?i)(d[ée]lai|lead.*time|delivery)", label):
                val = supplier.get("lead_time_days")
                if val:
                    cell.value = f"{val} jours"
                    if supplier.get("lead_time_source"):
                        from openpyxl.comments import Comment

                        cell.comment = Comment(supplier["lead_time_source"], "DMS")
                else:
                    cell.value = REVUE_MANUELLE
                    cell.fill = ORANGE_FILL

            elif re.search(r"(?i)(validit[ée]|validity)", label):
                val = supplier.get("validity_days")
                if val:
                    cell.value = f"{val} jours"
                    if supplier.get("validity_source"):
                        from openpyxl.comments import Comment

                        cell.comment = Comment(supplier["validity_source"], "DMS")
                else:
                    cell.value = REVUE_MANUELLE
                    cell.fill = ORANGE_FILL

            elif re.search(r"(?i)(r[ée]f[ée]rence|experience)", label):
                refs = supplier.get("technical_refs", [])
                if refs:
                    cell.value = ", ".join(refs[:3])
                else:
                    cell.value = REVUE_MANUELLE
                    cell.fill = ORANGE_FILL

    # ❌ SUPPRIMER onglets debug de l'export final
    # On ne crée PLUS le DMS_SUMMARY dans l'export (logs uniquement)
    debug_sheets = ["DMS_SUMMARY", "DEBUG", "TEMP", "SCRATCH", "NOTES"]
    for sheet_name in list(wb.sheetnames):
        for debug_pattern in debug_sheets:
            if debug_pattern in sheet_name.upper():
                # Supprimer complètement (pas masquer)
                wb.remove(wb[sheet_name])
                break

    # Save output
    out_dir = OUTPUTS_DIR / case_id
    out_dir.mkdir(exist_ok=True)
    out_name = f"CBA_{Path(template_path).stem}_{uuid.uuid4().hex[:6]}.xlsx"
    out_path = out_dir / out_name
    wb.save(out_path)

    return str(out_path)


# =========================
# PV Generation (Template-specific)
# =========================
def generate_pv_adaptive(
    template_path: str | None,
    case_id: str,
    case_title: str,
    suppliers: list[dict],
    dao_criteria: list[DAOCriterion],
    decision: dict | None = None,
) -> str:
    """
    Generate PV from template or create structured fallback.
    1 DAO = 1 PV template = 1 PV unique.
    """
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
    else:
        doc = Document()
        doc.add_heading("Procès-Verbal — Decision Memory System", level=1)
        doc.add_paragraph(f"Case: {case_title}")
        doc.add_paragraph(f"Case ID: {case_id}")
        doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}")

    # Suppliers summary
    lines: list[str] = []
    for s in suppliers:
        missing = s.get("missing_fields", [])
        status = "⚠️ Données incomplètes" if missing else "✓ Complet"
        lines.append(
            f"- {s['supplier_name']} | Prix: {s.get('total_price') or 'N/A'} | "
            f"Délai: {s.get('lead_time_days') or 'N/A'}j | "
            f"Validité: {s.get('validity_days') or 'N/A'}j | "
            f"Status: {status}"
        )
    suppliers_block = "\n".join(lines)

    # Placeholders replacement
    repl = {
        "{{CASE_ID}}": case_id,
        "{{CASE_TITLE}}": case_title,
        "{{GENERATED_AT}}": datetime.utcnow().isoformat(),
        "{{SUPPLIERS_TABLE}}": suppliers_block,
        "{{CRITERIA_COUNT}}": str(len(dao_criteria)),
        "{{OFFERS_COUNT}}": str(len(suppliers)),
        "{{CHOSEN_SUPPLIER}}": (
            decision["chosen_supplier"]
            if decision
            else "NON DÉCIDÉ (validation humaine requise)"
        ),
        "{{DECISION_REASON}}": (
            decision.get("decision_reason", "") if decision else ""
        ),
        "{{NEXT_ACTION}}": (decision.get("next_action", "") if decision else ""),
    }

    for p in doc.paragraphs:
        for k, v in repl.items():
            if k in p.text:
                p.text = p.text.replace(k, v)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for k, v in repl.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)

    # Append structured content if no template or empty
    if not template_path or doc.paragraphs[-1].text.strip() == "":
        doc.add_page_break()
        doc.add_heading("RÉSUMÉ AUTO-EXTRAIT (à valider)", level=2)

        doc.add_heading("Critères DAO", level=3)
        for c in dao_criteria:
            doc.add_paragraph(
                f"• {c.critere_nom} ({c.categorie}) — "
                f"{'Éliminatoire' if c.ponderation == 0 else f'{c.ponderation*100:.0f}%'}",
                style="List Bullet",
            )

        doc.add_heading("Offres reçues", level=3)
        for line in lines:
            doc.add_paragraph(line, style="List Bullet")

        if decision:
            doc.add_page_break()
            doc.add_heading("DÉCISION HUMAINE (FINALE)", level=2)
            doc.add_paragraph(f"Fournisseur retenu: {decision['chosen_supplier']}")
            doc.add_paragraph(f"Justification: {decision['decision_reason']}")
            doc.add_paragraph(f"Action suivante: {decision['next_action']}")
            doc.add_paragraph(f"Horodatage: {decision['decided_at']}")
            doc.add_paragraph(
                "\n⚠️ Cette décision est prise par le comité d'évaluation. "
                "Le système n'a effectué aucune recommandation automatique."
            )

    # Save output
    out_dir = OUTPUTS_DIR / case_id
    out_dir.mkdir(exist_ok=True)
    tpl_name = Path(template_path).stem if template_path else "DEFAULT"
    out_name = f"PV_{tpl_name}_{uuid.uuid4().hex[:6]}.docx"
    out_path = out_dir / out_name
    doc.save(out_path)

    return str(out_path)
