"""
Text extraction from different document formats (PDF, DOCX).
"""

from pathlib import Path
from typing import List

from docx import Document
from pypdf import PdfReader
from fastapi import HTTPException


# =========================
# Text Extraction
# =========================
def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    out: List[str] = []
    for i, page in enumerate(reader.pages):
        txt = (page.extract_text() or "").strip()
        if txt:
            out.append(f"[PAGE {i+1}]\n{txt}\n")
    return "\n".join(out).strip()


def extract_text_any(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(path)
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    raise HTTPException(status_code=400, detail=f"Unsupported extraction: {ext}")
