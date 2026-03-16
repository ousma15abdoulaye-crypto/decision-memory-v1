"""PV generation services for Couche A."""

from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any

from docx import Document

from ..models import (
    ensure_schema,
    generate_id,
    get_connection,
    serialize_json,
)


def _build_pv_doc(kind: str, lot_name: str, offers: list[dict]) -> Document:
    doc = Document()
    doc.add_heading(f"PV {kind}", level=1)
    doc.add_paragraph(f"Lot: {lot_name}")
    doc.add_paragraph("Synthèse des offres:")
    for offer in offers:
        doc.add_paragraph(
            f"- {offer['supplier_name']} | Montant: {offer.get('amount') or 'N/A'} | Statut: {offer.get('status')}"
        )
    return doc


async def generate_pv_ouverture(
    lot_id: str, actor: str | None = None
) -> dict[str, Any]:
    """Generate a PV d'ouverture for a lot."""
    return await _generate_pv(lot_id, "OUVERTURE", actor)


async def generate_pv_analyse(lot_id: str, actor: str | None = None) -> dict[str, Any]:
    """Generate a PV d'analyse for a lot."""
    return await _generate_pv(lot_id, "ANALYSE", actor)


async def _generate_pv(lot_id: str, kind: str, actor: str | None) -> dict[str, Any]:
    def _process() -> dict[str, Any]:
        ensure_schema()
        with get_connection() as conn:
            conn.execute(
                "SELECT * FROM lots WHERE id = %(lot_id)s",
                {"lot_id": lot_id},
            )
            lot = conn.fetchone()
            if not lot:
                raise ValueError("Lot introuvable.")

            conn.execute(
                """
                SELECT o.id, o.supplier_name, o.amount, o.lot_id, o.case_id,
                       a.status AS analysis_status
                FROM offers o
                LEFT JOIN analyses a ON a.offer_id = o.id
                WHERE o.lot_id = %(lot_id)s
                """,
                {"lot_id": lot_id},
            )
            offers = conn.fetchall()
            if not offers:
                raise ValueError("Aucune offre disponible pour ce lot.")

            lot_name = lot.get("name") or lot.get("lot_number") or lot_id
            offer_payload = [
                {
                    "supplier_name": o["supplier_name"],
                    "amount": o.get("amount"),
                    "status": o.get("analysis_status")
                    or o.get("status")
                    or "EN_ATTENTE",
                }
                for o in offers
            ]
            doc = _build_pv_doc(kind, lot_name, offer_payload)

            pv_dir = Path("data") / "couche_a" / "pv" / lot_id
            pv_dir.mkdir(parents=True, exist_ok=True)
            filename = f"pv_{kind.lower()}_{lot_id}.docx"
            pv_path = pv_dir / filename
            doc.save(pv_path)

            doc_id = generate_id()
            conn.execute(
                """
                INSERT INTO documents (
                    id, case_id, lot_id, offer_id, document_type, filename,
                    storage_path, path, mime_type, metadata_json, uploaded_at
                )
                VALUES (
                    %(id)s, %(case_id)s, %(lot_id)s, %(offer_id)s, %(document_type)s, %(filename)s,
                    %(storage_path)s, %(path)s, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    %(metadata_json)s, NOW()::TEXT
                )
                """,
                {
                    "id": doc_id,
                    "case_id": lot["case_id"],
                    "lot_id": lot_id,
                    "offer_id": offers[0]["id"],
                    "document_type": f"PV_{kind}",
                    "filename": filename,
                    "storage_path": str(pv_path),
                    "path": str(pv_path),
                    "metadata_json": serialize_json({"lot_id": lot_id}),
                },
            )
            conn.execute(
                """
                INSERT INTO audits (
                    id, case_id, entity_type, entity_id, action, actor, details_json, created_at
                )
                VALUES (
                    %(id)s, %(case_id)s, 'pv', %(entity_id)s, %(action)s, %(actor)s, %(details_json)s, NOW()::TEXT
                )
                """,
                {
                    "id": generate_id(),
                    "case_id": lot["case_id"],
                    "entity_id": doc_id,
                    "action": f"PV_{kind}",
                    "actor": actor,
                    "details_json": serialize_json({"lot_id": lot_id}),
                },
            )

        return {"document_id": doc_id, "path": str(pv_path)}

    return await asyncio.to_thread(_process)
