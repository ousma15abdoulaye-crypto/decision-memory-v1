import json
import os
import re
import uuid

# Load .env before db import (DATABASE_URL required)
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Configure logging pour resilience patterns
from src.logging_config import configure_logging
configure_logging()

from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict

from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from openpyxl.utils import get_column_letter

from pypdf import PdfReader

from src.db import get_connection, db_execute, db_execute_one, db_fetchall, init_db_schema
from src.couche_a.routers import router as upload_router
from src.auth_router import router as auth_router
from src.ratelimit import init_rate_limit, limiter
from src.auth import CurrentUser
from src.core.config import (
    APP_TITLE, APP_VERSION, BASE_DIR, DATA_DIR, UPLOADS_DIR, OUTPUTS_DIR, 
    STATIC_DIR, INVARIANTS
)
from src.core.models import (
    CaseCreate, AnalyzeRequest, DecideRequest,
    CBATemplateSchema, DAOCriterion, OfferSubtype, SupplierPackage
)
from src.core.dependencies import (
    safe_save_upload, register_artifact, get_artifacts,
    add_memory, list_memory
)
from src.business.extraction import extract_text_from_docx, extract_text_from_pdf, extract_text_any
from src.business.offer_processor import (
    detect_offer_subtype, aggregate_supplier_packages,
    guess_supplier_name, extract_offer_data_guided
)
from src.api import health, cases

# ❌ REMOVED: from src.couche_a.procurement import router as procurement_router (M2-Extended)

# =========================
# Database — PostgreSQL only (schema created on startup)
# =========================
from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app):
    init_db_schema()
    yield

app = FastAPI(title=APP_TITLE, version=APP_VERSION, lifespan=lifespan)

# Initialize rate limiting
init_rate_limit(app)

# Include routers
app.include_router(auth_router)
app.include_router(upload_router)
app.include_router(health.router)
app.include_router(cases.router)
# ❌ REMOVED: app.include_router(procurement_router) (M2-Extended)

app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")




# =========================
# CBA Template Analysis (ADAPTIVE)
# =========================
def analyze_cba_template(template_path: str) -> CBATemplateSchema:
    """Analyse dynamique structure template CBA"""
    wb = load_workbook(template_path)
    ws = wb.active

    supplier_header_row = None
    supplier_name_row = None
    supplier_cols: List[int] = []

    # Detect supplier header
    for row_idx in range(1, min(12, ws.max_row + 1)):
        for col_idx in range(1, ws.max_column + 1):
            val = ws.cell(row_idx, col_idx).value
            if val and isinstance(val, str):
                if re.search(r"(?i)(soumissionnaire|supplier|fournisseur|bidder)", val):
                    supplier_header_row = row_idx
                    # Detect supplier columns
                    for c in range(col_idx, min(col_idx + 15, ws.max_column + 1)):
                        if ws.cell(row_idx, c).value:
                            supplier_cols.append(c)
                    break
        if supplier_header_row:
            break

    supplier_name_row = supplier_header_row + 1 if supplier_header_row else None

    # Detect criteria rows
    criteria_start_row = None
    criteria_rows: List[Dict[str, Any]] = []

    if supplier_name_row:
        for row_idx in range(supplier_name_row + 1, min(supplier_name_row + 40, ws.max_row + 1)):
            col_a = ws.cell(row_idx, 1).value or ""
            col_b = ws.cell(row_idx, 2).value or ""

            if isinstance(col_b, str) and len(col_b) > 5:
                if not criteria_start_row:
                    criteria_start_row = row_idx

                # Type detection
                ctype = "autre"
                if re.search(r"(?i)(essent|mandatory|obligatoire)", col_b):
                    ctype = "essentiels"
                elif re.search(r"(?i)(technique|technical|capacity)", col_b):
                    ctype = "technique"
                elif re.search(r"(?i)(commercial|price|prix|co[uû]t)", col_b):
                    ctype = "commercial"
                elif re.search(r"(?i)(durabilit[ée]|sustainability)", col_b):
                    ctype = "durabilite"

                criteria_rows.append({
                    "row": row_idx,
                    "name": col_b[:150],
                    "type": ctype
                })

    schema = CBATemplateSchema(
        template_id=str(uuid.uuid4()),
        template_name=Path(template_path).name,
        supplier_header_row=supplier_header_row or 0,
        supplier_name_row=supplier_name_row or 0,
        supplier_cols=supplier_cols,
        criteria_start_row=criteria_start_row or 0,
        criteria_rows=criteria_rows,
        sheets=wb.sheetnames,
        meta={"detected_at": datetime.utcnow().isoformat()}
    )

    return schema


def fill_cba_adaptive(
    template_path: str,
    case_id: str,
    suppliers: List[dict],
    dao_criteria: List[DAOCriterion]
) -> str:
    """
    Remplissage adaptatif basé sur structure template détectée.
    1 DAO = 1 template spécifique = 1 CBA unique.
    
    COMPORTEMENT ATTENDU:
    - Noms fournisseurs RÉELS (pas d'IDs, pas de hash)
    - Données manquantes → "REVUE MANUELLE" avec surlignage ORANGE
    - Aucun onglet debug visible dans l'export final
    - Aucune note magique, aucune élimination implicite
    """
    schema = analyze_cba_template(template_path)

    # Save schema to DB for future learning
    with get_connection() as conn:
        db_execute(conn, """
            INSERT INTO cba_template_schemas (id, case_id, template_name, structure_json, created_at)
            VALUES (:tid, :cid, :tname, :struct, :ts)
        """, {
            "tid": schema.template_id, "cid": case_id, "tname": schema.template_name,
            "struct": json.dumps(asdict(schema), ensure_ascii=False),
            "ts": datetime.utcnow().isoformat(),
        })

    # Load template
    wb = load_workbook(template_path)
    ws = wb.active
    
    # Couleur ORANGE pour REVUE MANUELLE (spec conforme)
    ORANGE_FILL = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
    REVUE_MANUELLE = "REVUE MANUELLE"

    # Metadata section (top)
    ws["A1"] = "Decision Memory System — CBA Adaptatif"
    ws["A2"] = f"Case ID: {case_id}"
    ws["A3"] = f"Generated: {datetime.utcnow().isoformat()}"
    ws["A4"] = f"Template: {schema.template_name}"

    # Fill supplier names (detected row/cols) - NOMS RÉELS uniquement
    if schema.supplier_name_row and schema.supplier_cols:
        for idx, supplier in enumerate(suppliers[:len(schema.supplier_cols)]):
            col = schema.supplier_cols[idx]
            supplier_name = supplier.get("supplier_name", "")
            
            # Vérifier que ce n'est pas un ID
            if not supplier_name or supplier_name in ["SUPPLIER_UNKNOWN", "FOURNISSEUR_INCONNU"]:
                supplier_name = REVUE_MANUELLE
            
            # Écriture unique de la cellule
            cell = ws.cell(schema.supplier_name_row, col)
            cell.value = supplier_name
            if supplier_name == REVUE_MANUELLE:
                cell.fill = ORANGE_FILL

    # Fill commercial data (detected criteria rows)
    for row_idx in range(schema.criteria_start_row, min(schema.criteria_start_row + 50, ws.max_row + 1)):
        label = ws.cell(row_idx, 2).value or ""

        for idx, supplier in enumerate(suppliers[:len(schema.supplier_cols)]):
            col = schema.supplier_cols[idx]
            cell = ws.cell(row_idx, col)
            
            # Vérifier le package_status du fournisseur
            package_status = supplier.get("package_status", "UNKNOWN")
            has_financial = supplier.get("has_financial", False)

            # Mapping guided by criteria
            if re.search(r"(?i)(prix|price|montant|co[uû]t)", label):
                if has_financial:
                    val = supplier.get("total_price")
                    if val:
                        cell.value = val
                        if supplier.get("total_price_source"):
                            from openpyxl.comments import Comment
                            cell.comment = Comment(supplier["total_price_source"], "DMS")
                    else:
                        cell.value = REVUE_MANUELLE
                        cell.fill = ORANGE_FILL
                else:
                    # Offre sans partie financière
                    cell.value = REVUE_MANUELLE
                    cell.fill = ORANGE_FILL

            elif re.search(r"(?i)(d[ée]lai|lead.*time|delivery)", label):
                val = supplier.get("lead_time_days")
                if val:
                    cell.value = f"{val} jours"
                    if supplier.get("lead_time_source"):
                        from openpyxl.comments import Comment
                        cell.comment = Comment(supplier["lead_time_source"], "DMS")
                else:
                    cell.value = REVUE_MANUELLE
                    cell.fill = ORANGE_FILL

            elif re.search(r"(?i)(validit[ée]|validity)", label):
                val = supplier.get("validity_days")
                if val:
                    cell.value = f"{val} jours"
                    if supplier.get("validity_source"):
                        from openpyxl.comments import Comment
                        cell.comment = Comment(supplier["validity_source"], "DMS")
                else:
                    cell.value = REVUE_MANUELLE
                    cell.fill = ORANGE_FILL

            elif re.search(r"(?i)(r[ée]f[ée]rence|experience)", label):
                refs = supplier.get("technical_refs", [])
                if refs:
                    cell.value = ", ".join(refs[:3])
                else:
                    cell.value = REVUE_MANUELLE
                    cell.fill = ORANGE_FILL

    # ❌ SUPPRIMER onglets debug de l'export final
    # On ne crée PLUS le DMS_SUMMARY dans l'export (logs uniquement)
    debug_sheets = ["DMS_SUMMARY", "DEBUG", "TEMP", "SCRATCH", "NOTES"]
    for sheet_name in list(wb.sheetnames):
        for debug_pattern in debug_sheets:
            if debug_pattern in sheet_name.upper():
                # Supprimer complètement (pas masquer)
                wb.remove(wb[sheet_name])
                break

    # Save output
    out_dir = OUTPUTS_DIR / case_id
    out_dir.mkdir(exist_ok=True)
    out_name = f"CBA_{Path(template_path).stem}_{uuid.uuid4().hex[:6]}.xlsx"
    out_path = out_dir / out_name
    wb.save(out_path)

    return str(out_path)


# =========================
# PV Generation (Template-specific)
# =========================
def generate_pv_adaptive(
    template_path: Optional[str],
    case_id: str,
    case_title: str,
    suppliers: List[dict],
    dao_criteria: List[DAOCriterion],
    decision: Optional[dict] = None
) -> str:
    """
    Generate PV from template or create structured fallback.
    1 DAO = 1 PV template = 1 PV unique.
    """
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
    else:
        doc = Document()
        doc.add_heading("Procès-Verbal — Decision Memory System", level=1)
        doc.add_paragraph(f"Case: {case_title}")
        doc.add_paragraph(f"Case ID: {case_id}")
        doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}")

    # Suppliers summary
    lines: List[str] = []
    for s in suppliers:
        missing = s.get("missing_fields", [])
        status = "⚠️ Données incomplètes" if missing else "✓ Complet"
        lines.append(
            f"- {s['supplier_name']} | Prix: {s.get('total_price') or 'N/A'} | "
            f"Délai: {s.get('lead_time_days') or 'N/A'}j | "
            f"Validité: {s.get('validity_days') or 'N/A'}j | "
            f"Status: {status}"
        )
    suppliers_block = "\n".join(lines)

    # Placeholders replacement
    repl = {
        "{{CASE_ID}}": case_id,
        "{{CASE_TITLE}}": case_title,
        "{{GENERATED_AT}}": datetime.utcnow().isoformat(),
        "{{SUPPLIERS_TABLE}}": suppliers_block,
        "{{CRITERIA_COUNT}}": str(len(dao_criteria)),
        "{{OFFERS_COUNT}}": str(len(suppliers)),
        "{{CHOSEN_SUPPLIER}}": (decision["chosen_supplier"] if decision else "NON DÉCIDÉ (validation humaine requise)"),
        "{{DECISION_REASON}}": (decision.get("decision_reason", "") if decision else ""),
        "{{NEXT_ACTION}}": (decision.get("next_action", "") if decision else ""),
    }

    for p in doc.paragraphs:
        for k, v in repl.items():
            if k in p.text:
                p.text = p.text.replace(k, v)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for k, v in repl.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)

    # Append structured content if no template or empty
    if not template_path or doc.paragraphs[-1].text.strip() == "":
        doc.add_page_break()
        doc.add_heading("RÉSUMÉ AUTO-EXTRAIT (à valider)", level=2)

        doc.add_heading("Critères DAO", level=3)
        for c in dao_criteria:
            doc.add_paragraph(
                f"• {c.critere_nom} ({c.categorie}) — "
                f"{'Éliminatoire' if c.ponderation == 0 else f'{c.ponderation*100:.0f}%'}",
                style="List Bullet"
            )

        doc.add_heading("Offres reçues", level=3)
        for line in lines:
            doc.add_paragraph(line, style="List Bullet")

        if decision:
            doc.add_page_break()
            doc.add_heading("DÉCISION HUMAINE (FINALE)", level=2)
            doc.add_paragraph(f"Fournisseur retenu: {decision['chosen_supplier']}")
            doc.add_paragraph(f"Justification: {decision['decision_reason']}")
            doc.add_paragraph(f"Action suivante: {decision['next_action']}")
            doc.add_paragraph(f"Horodatage: {decision['decided_at']}")
            doc.add_paragraph("\n⚠️ Cette décision est prise par le comité d'évaluation. "
                            "Le système n'a effectué aucune recommandation automatique.")

    # Save output
    out_dir = OUTPUTS_DIR / case_id
    out_dir.mkdir(exist_ok=True)
    tpl_name = Path(template_path).stem if template_path else "DEFAULT"
    out_name = f"PV_{tpl_name}_{uuid.uuid4().hex[:6]}.docx"
    out_path = out_dir / out_name
    doc.save(out_path)

    return str(out_path)


# =========================
# API Routes
# =========================
@app.get("/", response_class=HTMLResponse)
def home():
    idx = STATIC_DIR / "index.html"
    if not idx.exists():
        return HTMLResponse("<h3>Missing static/index.html</h3>", status_code=500)
    return HTMLResponse(idx.read_text(encoding="utf-8"))


@app.get("/api/health")
def health():
    return {
        "status": "healthy",
        "version": APP_VERSION,
        "invariants_status": "enforced"
    }


@app.get("/api/constitution")
def api_constitution():
    return {"invariants": INVARIANTS, "version": APP_VERSION}


@app.post("/api/cases")
@limiter.limit("10/minute")
async def create_case(request: Request, payload: CaseCreate, user: CurrentUser):
    """Crée nouveau case (requiert authentification)."""
    case_type = payload.case_type.strip().upper()
    if case_type not in {"DAO", "RFQ"}:
        raise HTTPException(status_code=400, detail="case_type must be DAO or RFQ")

    case_id = str(uuid.uuid4())
    now = datetime.utcnow().isoformat()

    with get_connection() as conn:
        db_execute(conn, """
            INSERT INTO cases (id, case_type, title, lot, created_at, status, owner_id)
            VALUES (:id, :ctype, :title, :lot, :ts, :status, :owner)
        """, {
            "id": case_id,
            "ctype": case_type,
            "title": payload.title.strip(),
            "lot": payload.lot,
            "ts": now,
            "status": "open",
            "owner": user["id"]
        })

    return {
        "id": case_id,
        "case_type": case_type,
        "title": payload.title,
        "lot": payload.lot,
        "created_at": now,
        "status": "open",
        "owner_id": user["id"]
    }


@app.get("/api/cases")
@limiter.limit("50/minute")
def list_cases(request: Request):
    """Liste tous les cases (rate limited)."""
    with get_connection() as conn:
        rows = db_fetchall(conn, "SELECT * FROM cases ORDER BY created_at DESC")
    return rows


@app.get("/api/cases/{case_id}")
def get_case(case_id: str):
    with get_connection() as conn:
        c = db_execute_one(conn, "SELECT * FROM cases WHERE id=:id", {"id": case_id})
    if not c:
        raise HTTPException(status_code=404, detail="case not found")

    arts = get_artifacts(case_id)
    mem = list_memory(case_id)

    # Get DAO criteria if analyzed
    with get_connection() as conn:
        criteria = db_fetchall(conn, "SELECT * FROM dao_criteria WHERE case_id=:cid ORDER BY ordre_affichage", {"cid": case_id})

    return {
        "case": dict(c),
        "artifacts": arts,
        "memory": mem,
        "dao_criteria": criteria
    }


@app.post("/api/upload/{case_id}/{kind}")
def upload(case_id: str, kind: str, file: UploadFile = File(...)):
    with get_connection() as conn:
        c = db_execute_one(conn, "SELECT id FROM cases WHERE id=:id", {"id": case_id})
    if not c:
        raise HTTPException(status_code=404, detail="case not found")

    kind = kind.strip().lower()
    allowed = {"dao", "offer", "cba_template", "pv_template"}
    if kind not in allowed:
        raise HTTPException(status_code=400, detail=f"kind must be one of {sorted(list(allowed))}")

    filename, path = safe_save_upload(case_id, kind, file)
    aid = register_artifact(case_id, kind, filename, path, meta={"original_name": file.filename})

    return {"ok": True, "artifact_id": aid, "filename": filename}


@app.post("/api/analyze")
def analyze(payload: AnalyzeRequest):
    """
    Main analysis pipeline:
    1. Extract DAO criteria (structured)
    2. Extract offer data (DAO-guided)
    3. Generate CBA (template-adaptive)
    4. Generate PV (template-specific)
    5. Store memory (passive)
    """
    case_id = payload.case_id

    with get_connection() as conn:
        case = db_execute_one(conn, "SELECT * FROM cases WHERE id=:id", {"id": case_id})
    if not case:
        raise HTTPException(status_code=404, detail="case not found")

    # Step 1: DAO extraction
    dao_arts = get_artifacts(case_id, "dao")
    if not dao_arts:
        raise HTTPException(status_code=400, detail="Missing DAO (upload kind=dao)")

    dao_path = dao_arts[0]["path"]
    dao_text = extract_text_any(dao_path)
    dao_criteria = extract_dao_criteria_structured(dao_text)

    # Store criteria in DB
    with get_connection() as conn:
        for c in dao_criteria:
            db_execute(conn, """
                INSERT INTO dao_criteria 
                (id, case_id, categorie, critere_nom, description, ponderation, type_reponse, seuil_elimination, ordre_affichage, created_at)
                VALUES (:id, :cid, :cat, :nom, :desc, :pond, :type_reponse, :seuil, :ordre, :ts)
            """, {
                "id": str(uuid.uuid4()), "cid": case_id, "cat": c.categorie, "nom": c.critere_nom,
                "desc": c.description, "pond": c.ponderation, "type_reponse": c.type_reponse,
                "seuil": c.seuil_elimination, "ordre": c.ordre_affichage,
                "ts": datetime.utcnow().isoformat(),
            })

    # Step 2: Offers extraction (DAO-guided) + SUBTYPE DETECTION
    offer_arts = get_artifacts(case_id, "offer")
    if not offer_arts:
        raise HTTPException(status_code=400, detail="Missing OFFERS (upload kind=offer)")

    raw_offers: List[dict] = []
    for off in offer_arts:
        txt = extract_text_any(off["path"])
        supplier_name = guess_supplier_name(txt, off["filename"])
        
        # CRITIQUE: Détection du subtype (FINANCIAL_ONLY, etc.)
        subtype = detect_offer_subtype(txt, off["filename"])
        
        offer_data = extract_offer_data_guided(txt, dao_criteria)

        raw_offers.append({
            "supplier_name": supplier_name,
            "subtype": asdict(subtype),
            **offer_data,
            "source_filename": off["filename"],
            "artifact_id": off["id"]
        })

        # Store extraction in DB
        with get_connection() as conn:
            db_execute(conn, """
                INSERT INTO offer_extractions (id, case_id, artifact_id, supplier_name, extracted_data_json, missing_fields_json, created_at)
                VALUES (:id, :cid, :aid, :supplier, :extracted, :missing, :ts)
            """, {
                "id": str(uuid.uuid4()), "cid": case_id, "aid": off["id"], "supplier": supplier_name,
                "extracted": json.dumps({**offer_data, "subtype": asdict(subtype)}, ensure_ascii=False),
                "missing": json.dumps(offer_data.get("missing_fields", []), ensure_ascii=False),
                "ts": datetime.utcnow().isoformat(),
            })

    # CRITIQUE: Agrégation par fournisseur (gestion offres partielles)
    supplier_packages = aggregate_supplier_packages(raw_offers)
    
    # Conversion en format compatible avec le reste du système
    suppliers: List[dict] = []
    for pkg in supplier_packages:
        suppliers.append({
            "supplier_name": pkg.supplier_name,
            "package_status": pkg.package_status,
            "has_financial": pkg.has_financial,
            "has_technical": pkg.has_technical,
            "has_admin": pkg.has_admin,
            **pkg.extracted_data,
            "offer_ids": pkg.offer_ids,
            "document_count": len(pkg.documents)
        })

    # Memory entry avec traçabilité des décisions
    add_memory(case_id, "extraction", {
        "dao_source": dao_arts[0]["filename"],
        "offers_sources": [o["filename"] for o in offer_arts],
        "dao_criteria_count": len(dao_criteria),
        "dao_criteria": [asdict(c) for c in dao_criteria],
        "raw_offers_count": len(raw_offers),
        "supplier_packages_count": len(supplier_packages),
        "packages_summary": [
            {
                "supplier": pkg.supplier_name,
                "status": pkg.package_status,
                "subtypes": [d.get("subtype", {}).get("subtype") for d in pkg.documents]
            }
            for pkg in supplier_packages
        ],
        "offers_summary": suppliers,
        "note": "DAO-driven extraction with partial offers support. Subtype detection enabled. No scoring/ranking. Human validates."
    })

    # Step 3: Generate CBA (adaptive)
    cba_out = None
    cba_tpl = get_artifacts(case_id, "cba_template")
    if cba_tpl:
        cba_out = fill_cba_adaptive(cba_tpl[0]["path"], case_id, suppliers, dao_criteria)
        register_artifact(case_id, "output_cba", Path(cba_out).name, cba_out, 
                        meta={"template_used": cba_tpl[0]["filename"]})

    # Step 4: Generate PV (adaptive)
    pv_tpl = get_artifacts(case_id, "pv_template")
    pv_tpl_path = pv_tpl[0]["path"] if pv_tpl else None
    pv_out = generate_pv_adaptive(pv_tpl_path, case_id, case["title"], suppliers, dao_criteria, decision=None)
    register_artifact(case_id, "output_pv", Path(pv_out).name, pv_out, 
                     meta={"template_used": pv_tpl[0]["filename"] if pv_tpl else "default", "decision_included": False})

    # Statistiques et warnings
    partial_offers = [s for s in suppliers if s.get("package_status") == "PARTIAL"]
    complete_offers = [s for s in suppliers if s.get("package_status") == "COMPLETE"]
    financial_only = [s for s in suppliers if s.get("has_financial") and not s.get("has_technical")]
    
    return {
        "ok": True,
        "case_id": case_id,
        "dao_criteria_count": len(dao_criteria),
        "offers_count": len(suppliers),
        "raw_documents_count": len(raw_offers),
        "output_cba_generated": bool(cba_out),
        "output_pv_generated": True,
        "downloads": {
            "cba": f"/api/download/{case_id}/output_cba" if cba_out else None,
            "pv": f"/api/download/{case_id}/output_pv",
        },
        "package_stats": {
            "complete": len(complete_offers),
            "partial": len(partial_offers),
            "financial_only": len(financial_only)
        },
        "warnings": {
            "missing_data_count": sum(len(s.get("missing_fields", [])) for s in suppliers),
            "suppliers_with_missing_data": [s["supplier_name"] for s in suppliers if s.get("missing_fields")],
            "partial_offers_detected": len(partial_offers) > 0,
            "note": "Offres partielles gérées en mode LENIENT. Aucune pénalité automatique. Champs manquants marqués REVUE MANUELLE."
        }
    }


@app.post("/api/decide")
def decide(payload: DecideRequest):
    """Record human decision and regenerate PV"""
    case_id = payload.case_id

    with get_connection() as conn:
        case = db_execute_one(conn, "SELECT * FROM cases WHERE id=:id", {"id": case_id})
    if not case:
        raise HTTPException(status_code=404, detail="case not found")

    decision = {
        "chosen_supplier": payload.chosen_supplier.strip(),
        "decision_reason": payload.decision_reason.strip(),
        "next_action": payload.next_action.strip(),
        "decided_at": datetime.utcnow().isoformat(),
        "human_final": True,
    }
    add_memory(case_id, "decision", decision)

    # Get latest extraction data
    with get_connection() as conn:
        extractions = db_fetchall(conn, "SELECT * FROM offer_extractions WHERE case_id=:cid", {"cid": case_id})
        criteria_rows = db_fetchall(conn, "SELECT * FROM dao_criteria WHERE case_id=:cid", {"cid": case_id})

    suppliers = []
    for ext in extractions:
        data = json.loads(ext["extracted_data_json"])
        suppliers.append({
            "supplier_name": ext["supplier_name"],
            **data,
            "artifact_id": ext["artifact_id"]
        })

    dao_criteria = [
        DAOCriterion(
            categorie=r["categorie"],
            critere_nom=r["critere_nom"],
            description=r["description"] or "",
            ponderation=r["ponderation"],
            type_reponse=r["type_reponse"],
            seuil_elimination=r["seuil_elimination"],
            ordre_affichage=r["ordre_affichage"]
        )
        for r in criteria_rows
    ]

    # Regenerate PV with decision
    pv_tpl = get_artifacts(case_id, "pv_template")
    pv_tpl_path = pv_tpl[0]["path"] if pv_tpl else None
    pv_out = generate_pv_adaptive(pv_tpl_path, case_id, case["title"], suppliers, dao_criteria, decision=decision)
    register_artifact(case_id, "output_pv", Path(pv_out).name, pv_out, 
                     meta={"template_used": pv_tpl[0]["filename"] if pv_tpl else "default", "decision_included": True})

    return {
        "ok": True,
        "case_id": case_id,
        "pv_with_decision": f"/api/download/{case_id}/output_pv"
    }


@app.get("/api/download/{case_id}/{kind}")
def download_latest(case_id: str, kind: str):
    """Download latest generated artifact"""
    kind = kind.strip().lower()
    if kind not in {"output_cba", "output_pv"}:
        raise HTTPException(status_code=400, detail="kind must be output_cba or output_pv")

    arts = get_artifacts(case_id, kind)
    if not arts:
        raise HTTPException(status_code=404, detail=f"No artifact found for {kind}")

    p = Path(arts[0]["path"])
    if not p.exists():
        raise HTTPException(status_code=404, detail="file missing on disk")

    return FileResponse(path=str(p), filename=p.name, media_type="application/octet-stream")


@app.get("/api/memory/{case_id}")
def memory(case_id: str):
    """Retrieve all memory entries for a case"""
    return {"case_id": case_id, "memory": list_memory(case_id)}


@app.get("/api/search_memory/{case_id}")
def search_memory(case_id: str, q: str):
    """Search memory entries by keyword"""
    q = (q or "").strip().lower()
    if not q:
        return {"case_id": case_id, "hits": []}

    mem = list_memory(case_id)
    hits = []
    for entry in mem:
        blob = json.dumps(entry.get("content", {}), ensure_ascii=False).lower()
        if q in blob:
            hits.append({
                "id": entry["id"],
                "entry_type": entry["entry_type"],
                "created_at": entry["created_at"],
                "preview": entry["content"]
            })

    return {"case_id": case_id, "q": q, "hits": hits}


# Rebuild Pydantic models to resolve forward references
try:
    CaseCreate.model_rebuild()
    AnalyzeRequest.model_rebuild()
    DecideRequest.model_rebuild()
except Exception:
    pass  # Ignore if models don't need rebuilding

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)