"""Couche A – PV (Procès-Verbal) document generation."""

from __future__ import annotations

import os
from datetime import datetime, timezone

from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.couche_a.models import Submission
from backend.system.settings import get_settings

__all__ = ["generate_pv_opening", "generate_pv_analysis"]


async def generate_pv_opening(case_id: str, lot_id: str, db: AsyncSession) -> str:
    """Generate a PV d'ouverture (opening minutes) as a DOCX file."""
    settings = get_settings()
    output_dir = os.path.join(settings.OUTPUT_DIR, case_id, lot_id)
    os.makedirs(output_dir, exist_ok=True)

    stmt = select(Submission).where(
        Submission.case_id == case_id, Submission.lot_id == lot_id
    )
    submissions = (await db.execute(stmt)).scalars().all()

    doc = Document()
    doc.add_heading("Procès-Verbal d'Ouverture des Plis", level=1)
    doc.add_paragraph(f"Affaire : {case_id}  |  Lot : {lot_id}")
    doc.add_paragraph(f"Date : {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")
    doc.add_heading("Soumissions reçues", level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Fournisseur"
    hdr[2].text = "Canal"
    hdr[3].text = "Date réception"

    for idx, sub in enumerate(submissions, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = sub.vendor_name
        row[2].text = sub.channel
        row[3].text = sub.created_at.strftime("%Y-%m-%d %H:%M") if sub.created_at else ""

    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(output_dir, f"PV_ouverture_{case_id}_{lot_id}_{ts}.docx")
    doc.save(file_path)
    return file_path


async def generate_pv_analysis(case_id: str, lot_id: str, db: AsyncSession) -> str:
    """Generate a PV d'analyse (analysis minutes) as a DOCX file."""
    settings = get_settings()
    output_dir = os.path.join(settings.OUTPUT_DIR, case_id, lot_id)
    os.makedirs(output_dir, exist_ok=True)

    from backend.couche_a.rules_engine import evaluate_submission

    stmt = select(Submission).where(
        Submission.case_id == case_id, Submission.lot_id == lot_id
    )
    submissions = (await db.execute(stmt)).scalars().all()

    doc = Document()
    doc.add_heading("Procès-Verbal d'Analyse des Offres", level=1)
    doc.add_paragraph(f"Affaire : {case_id}  |  Lot : {lot_id}")
    doc.add_paragraph(f"Date : {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")
    doc.add_heading("Résultats de l'évaluation", level=2)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Fournisseur"
    hdr[1].text = "Capacité"
    hdr[2].text = "Durabilité"
    hdr[3].text = "Commercial"
    hdr[4].text = "Total"
    hdr[5].text = "Statut"

    for sub in submissions:
        ev = await evaluate_submission(sub.id, db)
        row = table.add_row().cells
        row[0].text = sub.vendor_name
        row[1].text = str(ev["scores"]["capacity"])
        row[2].text = str(ev["scores"]["durability"])
        row[3].text = str(ev["scores"]["commercial"])
        row[4].text = str(ev["total"])
        row[5].text = ev["status"]

    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(output_dir, f"PV_analyse_{case_id}_{lot_id}_{ts}.docx")
    doc.save(file_path)
    return file_path
